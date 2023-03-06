import math
import os
from datetime import datetime
import itertools

#For image type carrier files encoding/decoding
import cv2
import numpy as np

#For word doc type carrier files encoding/decoding
from docx import Document

PROGRAM_CONFIG_CONSOLEOUTPUT = 1
PROGRAM_CONFIG_DOENCODING = False
PROGRAM_CONFIG_DODECODING = False

PROGRAM_CONFIG_FRAMEHEADEROFFSET = 256

#===========================================================================
#Program-tested working carrier formats:
#	Doc: txt
#	Img: png, jpg, jpeg, bmp
#	Aud: acc, ac3, aiff, au, flac, m4a, m4b, m4r, mka, mp3, wav
#	Vid: m4v, mkv, mp4 [H.264], mpg, vob
#---------------------------------------------------------------------------
# **NOTE: 'jpg', jpeg' & 'jpe' are only supported by encoder through 
# conversion to png after encoding to bypass lossy compression. They 
# are not supported by decoder.
#===========================================================================
#Program-tested non-working carrier formats:
#	Doc: 
#	Img: 
#	Aud: ogg, wma
#	Vid: avi, asf, mp4 [H.265], wmv
#---------------------------------------------------------------------------
# Uses lossy compressions; alters encoded payload data:
# 	ogg, wma, wmv
# Highly compressed; corrupts encoded carrier:
#	asf, mp4 [H.265]
#===========================================================================

PROGRAM_CONFIG_ENCODEDCARRIERFORMAT_BLACKLIST = [
	"jpg", "jpeg", "jpe",
	"ogg", "wma",
	"asf", "wmv"
]

PROGRAM_CONFIG_CARRIERFORMAT_SPECIALLIST = {
	"images": ["png", "jpg", "jpeg"],
	"worddoc": ["docx", "doc"]
}

fileFormat = "bmp"

#Encoding Methods:
#	00: 2 Bit X-OR onion complier
#	01: Bit replacement
ENCODING_CONFIG_ENCODINGMETHOD = 1
ENCODING_CONFIG_SIGNATURES = {
	0: [ 239, 138,  95,  60,  95, 138, 239],
	1: [ 117,  95,  89,  60,  89,  95, 117]
}

ENCODING_INPUTFILE_PAYLOAD = "\Input\Images\TestFile01.bmp"
ENCODING_INPUTFILE_CARRIER = "\Input\Images\TestFile." + str(fileFormat)
ENCODING_OUTPUTFILE_ENCODEDCARRIER = "EncodedCarrier"
ENCODING_BITREPLACEMENT_HEX = 1
ENCODING_BITREPLACEMENT_WORDDOC_RESERVED = 126

DECODING_INPUTFILE_ENCODEDCARRIER = "\Output\EncodedCarrier." + str(fileFormat)
DECODING_OUTPUTFOLDER = "\Output"

#Main
def main():
	"""Main program directives."""

	printT("=== Program starting ===", True)

	#Acquire runtime root path
	dynamicRootPath = os.getcwd()
	encodedCarrierFilePath = (
		dynamicRootPath + 
		"\\Output\\" + 
		ENCODING_OUTPUTFILE_ENCODEDCARRIER + 
		"." + 
		ENCODING_INPUTFILE_CARRIER.split('.')[-1])

	printT("Main: Program initialization complete.")

	#Start encoding
	if(PROGRAM_CONFIG_DOENCODING):
		printT("Main: Starting encoder...")
		status, bitsModified = payloadToCarrierEncoder(
			dynamicRootPath + ENCODING_INPUTFILE_PAYLOAD,
			dynamicRootPath + ENCODING_INPUTFILE_CARRIER,
			encodedCarrierFilePath,
			ENCODING_CONFIG_ENCODINGMETHOD
		)
		if(status == 0):
			printT("Main: " + str(bitsModified) + " bits modified.")
			printT("Main: Encoding complete.")
		elif(status == 1):
			printT("*ERROR* Main: Encoder encountered an error.")
		elif(status == 2):
			printT("*ERROR* Main: Unsupport carrier file format detect.")
		else:
			printT("*ERROR* Main: Encoder encountered an unexpected error.")

	#Start decoding
	if(PROGRAM_CONFIG_DODECODING):
		printT("Main: Starting decoder...")

		status, decoderMsg = payloadFromCarrierDecoder(
			dynamicRootPath + DECODING_INPUTFILE_ENCODEDCARRIER,
			dynamicRootPath + DECODING_OUTPUTFOLDER
		)
		if(status == 0):
			printT("Main: '" + str(decoderMsg) + "' extracted.")
			printT("Main: Decoding complete.")
		else:
			printT("*ERROR* Main: Decoder failed.")

	#EOF
	printT("=== Program complete ===", True)

#Encoder/Decoder manager(s)
#	Default APIs for encoding & decoding, carries out validations
#	before delegating to suitable encoding/decoding 
# 	function(s)
def payloadToCarrierEncoder(
	payloadFilePath,
	carrierFilePath,
	outputFilePath,
	encodingMethod,
	encodingParam = ENCODING_BITREPLACEMENT_HEX,
	carrierHeaderSize = PROGRAM_CONFIG_FRAMEHEADEROFFSET,
	carrierFormatSpecialList = PROGRAM_CONFIG_CARRIERFORMAT_SPECIALLIST):
	"""Manages process of encoding payload into carrier."""

	printT("Encoder: Encoder commencing...")

	#Load payload file
	printT("Encoder: Loading payload file...")
	payloadBytes = bytearray()
	try:
		with open(payloadFilePath, "rb") as byteReader:
			currentByte = byteReader.read(2048)
			while currentByte:
				payloadBytes += currentByte
				currentByte = byteReader.read(2048)
		printT("Encoder: Payload file loaded.")
		printT("Encoder: Payload byte count: " + str(len(payloadBytes)))
	except FileNotFoundError:
		printT("*ERROR* Encoder: Invalid file path.")
		return 1, None
	except:
		printT("*ERROR* Encoder: Unexpected reading error occured.")
		return 1, None

	#Add encoding metadata to payload
	status, payloadBytes = payloadToCarrierEncoder_payloadEncodingMetadataPrepender(
		payloadBytes,
		os.path.basename(payloadFilePath))
	printT("Encoder: Encoding metadata added.")

	#Initialize runtime variable(s)
	carrierFileFormat = (carrierFilePath.split(".")[-1]).lower()
	status = 0
	carrierBytes = bytearray()
	encodedCarrierBytes = bytearray()
	imageDimensions = []

	#Initialize return variable(s)
	bitsModified = 0

	#Data reading and transforming for image type carrier files
	printT("Encoder: Loading carrier file...")
	if(carrierFileFormat in carrierFormatSpecialList.get("images")):
		printT("Encoder: Image carrier format detected.")

		#Load carrier file
		try:
			printT("Encoder: Loading carrier file...")
			carrierBytes = cv2.imread(carrierFilePath)
			printT("Encoder: Carrier file loaded.")
		except FileNotFoundError:
			printT("*ERROR* Encoder: Invalid file path.")
			return 1, None
		except:
			printT("*ERROR* Encoder: Unexpected reading error occured.")
			return 1, None
		
		#Store pixel data orginal dimensions
		imageDimensions = carrierBytes.shape

		#Flatten pixel data to single dimensional bytearray
		status, carrierBytes = multiDimensionArrayFlattener(carrierBytes)
		if(status == 1):
			printT("*ERROR* Encoder: Unexpected error occured while flattening encoded carrier bytes.")
			return 1, None
		printT("Encoder: Carrier data transformed.")
	#Data reading and transforming for word doc type carrier files
	elif(carrierFileFormat in carrierFormatSpecialList.get("worddoc")):
		printT("Encoder: Word document carrier format detected.")

		#Load carrier file
		try:
			#Read word document file
			carrierFile = Document(carrierFilePath)
			#Convert word document text to bytearray
			for paragraph in carrierFile.paragraphs:
				carrierBytes += bytearray(paragraph.text, "utf-8")
		except FileNotFoundError:
			printT("*ERROR* Encoder: Invalid file path.")
			return 1, None
		except:
			printT("*ERROR* Encoder: Unexpected reading error occured.")
			return 1, None
	#Data reading for non-specific type carrier files
	else:
		#Load carrier file
		try:
			with open(carrierFilePath, "rb") as byteReader:
				currentByte = byteReader.read(2048)
				while currentByte:
					carrierBytes += currentByte
					currentByte = byteReader.read(2048)
		except FileNotFoundError:
			printT("*ERROR* Encoder: Invalid file path.")
			return 1, None
		except:
			printT("*ERROR* Encoder: Unexpected reading error occured.")
			return 1, None
	printT("Encoder: Loading carrier file complete.")
	printT("Encoder: Carrier byte count: " + str(len(carrierBytes)))

	#Header processing for image & word doc type carrier files
	if((carrierFileFormat in carrierFormatSpecialList.get("images")) or 
	(carrierFileFormat in carrierFormatSpecialList.get("worddoc"))):
		#Use starting bytes of file content as header
		carrierEncodingSignatureBytes = carrierBytes[:16]
		carrierBytes = carrierBytes[16:]
	#Header processing for non-specific type carrier files
	else:
		carrierEncodingSignatureBytes = carrierBytes[carrierHeaderSize:(carrierHeaderSize + 16)]
		carrierHeaderBytes = carrierBytes[:carrierHeaderSize]
		carrierBytes = carrierBytes[(carrierHeaderSize + 16):]
	printT("Encoder: Header processing complete.")

	#Encode encoding method's signature into carrierHeaderBytes last 16 bytes
	status, carrierEncodingSignatureBytes = payloadToCarrierEncoder_encodingSignatureEncoder(
		carrierEncodingSignatureBytes, encodingMethod, encodingParam
	)
	#Error correction for word doc type carrier files
	if(carrierFileFormat in carrierFormatSpecialList.get("worddoc")):
		status, carrierEncodingSignatureBytes = twoBitXOrOnionEncoder_wordDocErrorCorrection(
			carrierEncodingSignatureBytes, 32, 126)
	if(status == 1):
		printT("*ERROR* Encoder: Failed to encode encoding method signature.")
		return 1, None

	#Store encoding signature method for image & word doc type carrier files
	if((carrierFileFormat in carrierFormatSpecialList.get("images")) or 
	(carrierFileFormat in carrierFormatSpecialList.get("worddoc"))):
		carrierHeaderBytes = carrierEncodingSignatureBytes
	#Include header frame offset bytes for non-specific type carrier files
	else:
		carrierHeaderBytes += carrierEncodingSignatureBytes
	printT("Encoder: Encoded encoding method signature.")

	#Start encoding logic
	printT("Encoder: Starting encoding...")
	#Encoding method 0
	if(encodingMethod == 0):
		printT("Encoder: Encoding method 0 selected.")
		status, encodedCarrierBytes, bitsModified = twoBitXOrOnionEncoder(
			carrierBytes, payloadBytes)

		#Error correction for word doc type carrier files
		if(carrierFileFormat in carrierFormatSpecialList.get("worddoc")):
			status, encodedCarrierBytes = twoBitXOrOnionEncoder_wordDocErrorCorrection(
				encodedCarrierBytes, 32, 126)
	#Encoding method 1
	elif(encodingMethod == 1):
		printT("Encoder: Encoding method 1 selected.")
		#Encoding method 02 for word doc type carrier files
		if(carrierFileFormat in carrierFormatSpecialList.get("worddoc")):
			status, encodedCarrierBytes, bitsModified = bitReplacementEncoder_wordDoc(
				carrierBytes, payloadBytes, encodingParam)
		else:
			status, encodedCarrierBytes, bitsModified = bitReplacementEncoder(
				carrierBytes, payloadBytes, encodingParam)

	#If encoding logic returned error status
	if(status == 1):
		printT("*ERROR* Encoder: Encoding method error detected.")
		return 1, None
	elif(status == 2):
		printT("*ERROR* Encoder: Carrier file too small.")
		return 1, None
	elif(status == 3):
		printT("*ERROR* Encoder: Invalid encoding parameter detected.")
		return 1, None
	else:
		printT("Encoder: Encoding successful.")

	#File writing logic for image type carrier files
	if(carrierFileFormat in carrierFormatSpecialList.get("images")):
		#Prepend encoding method signature
		encodedCarrierBytes = carrierHeaderBytes + encodedCarrierBytes

		#Reshape bytearray to pixel data
		status, encodedCarrierBytes = multiDimensionArrayReshaper(
			encodedCarrierBytes,
			imageDimensions)
		if(status == 1):
			printT("*ERROR* Encoder: Unexpected error occured while reshaping encoded carrier bytes.")
			return 1, None
		encodedCarrierBytes = np.asarray(encodedCarrierBytes)
		printT("Encoder: Encoded carrier data reshaped.")

		#Specify compression and resolution arguments for image type 
		#carrier files
		if(carrierFileFormat in ["jpg", "jpeg"]):
			printT("Encoder: Lossy image file type detected, using default format: png")
			outputFilePath = "".join(outputFilePath.split(".")[:-1])
			outputFilePath += ".png"

		#Start writing encoded data
		printT("Encoder: Writing encoded carrier file...")
		try:
			if os.path.exists(outputFilePath):
				os.remove(outputFilePath)
			cv2.imwrite(outputFilePath, encodedCarrierBytes)
		except:
			printT("*ERROR* Encoder: Unexpected writing error occured.")
			return 1, None

		return 0, bitsModified
	#File writing logic for word doc type carrier files
	elif(carrierFileFormat in carrierFormatSpecialList.get("worddoc")):
		printT("Encoder: Writing encoded carrier file...")
		try:
			encodedCarrierBytes = carrierHeaderBytes + encodedCarrierBytes
			encodedCarrierText = encodedCarrierBytes.decode("utf-8")

			encodedDoc = Document()
			paragraph = encodedDoc.add_paragraph().add_run(encodedCarrierText)
			paragraph.font.name = "Arial"

			if os.path.exists(outputFilePath):
				os.remove(outputFilePath)
			encodedDoc.save(outputFilePath)
		except:
			printT("*ERROR* Encoder: Unexpected writing error occured.")
			return 1, None
		return 0, bitsModified
	#File writing logic for non-specific type carrier files
	else:
		#Start writing encoded bytes
		printT("Encoder: Writing encoded carrier file...")
		encodedCarrierByteCount = len(carrierHeaderBytes) + len(encodedCarrierBytes)
		try:
			if os.path.exists(outputFilePath):
				os.remove(outputFilePath)
			with open(outputFilePath, "wb") as byteWriter:
				byteWriter.write(carrierHeaderBytes)
				byteWriter.write(encodedCarrierBytes)
		except:
			printT("*ERROR* Encoder: Unexpected writing error occured.")
			return 1, None
	printT("Encoder: Writing carrier file complete.")
		
	printT("Encoder: Encoded carrier byte count: " + str(encodedCarrierByteCount))
	printT("Encoder: Encoded carrier body byte count: " + str(len(encodedCarrierBytes)))
	printT("Encoder: Encoder terminating.")

	return 0, bitsModified
def payloadToCarrierEncoder_payloadEncodingMetadataPrepender(
	payloadBytes,
	payloadFileName):
	"""payloadToCarrierEncoder's sub function; helps prepend 
	encoding meta data to payload's bytearray."""
	
	#Acquire payload file's name
	printT("Encoder: Payload file name detected: '" + str(payloadFileName) + "'")

	#Prepare payload header segment01[1 byte]: Payload file name byte count
	payloadHeaderBytes = bytearray()
	payloadFileNameBytes = bytearray(payloadFileName, "utf-8")
	payloadFileNameByteCount = len(payloadFileNameBytes)
	payloadHeaderBytes += payloadFileNameByteCount.to_bytes(1, 'big')
	printT("Encoder: Payload name byte count: " + str(payloadFileNameByteCount))

	#Prepare payload header segment02: Payload file name bytes
	payloadHeaderBytes += payloadFileNameBytes

	#Prepare payload header segment03[4 bytes]: Payload body byte count
	payloadByteCount = len(payloadBytes)
	payloadByteCountBytes = payloadByteCount.to_bytes(4, 'big')
	payloadHeaderBytes += payloadByteCountBytes
	printT("Encoder: Payload byte count: " + str(payloadByteCount))

	#Prepend payload header to payload
	payloadBytes = payloadHeaderBytes + payloadBytes

	return 0, payloadBytes
def payloadToCarrierEncoder_encodingSignatureEncoder(
	carrierSignatureBytes,
	encodingMethod,
	encodingParam = 0
	):
	"""payloadToCarrierEncoder's sub function; helps encode encoding 
	method's signature using a fixed encoding method 16 carrier bytes."""

	#Acquire encoding signature through dictionary
	encodingSignature = (ENCODING_CONFIG_SIGNATURES.get(encodingMethod)).copy()

	#Append 1 byte of parameters for encoding methods
	encodingSignature.append(encodingParam)

	#Initializing runtime variables
	encodedEncodingSignatureBytes = carrierSignatureBytes.copy()
	signatureBytes = bytearray()

	#Convert encoding signature to bytearray
	for x in encodingSignature:
		signatureBytes += x.to_bytes(1, 'big')

	#Encode signature bytearray into last 16 bytes of carrier header
	status, encodedEncodingSignatureBytes, bitsModified = twoBitXOrOnionEncoder(
		carrierSignatureBytes, signatureBytes)

	if(status == 1):
		return status, None
	else:
		return status, encodedEncodingSignatureBytes
def payloadFromCarrierDecoder(
	encodedCarrierFilePath,
	decodedPayloadDestPath,
	carrierHeaderSize = PROGRAM_CONFIG_FRAMEHEADEROFFSET,
	carrierFormatBlacklist = PROGRAM_CONFIG_ENCODEDCARRIERFORMAT_BLACKLIST,
	carrierFormatSpecialList = PROGRAM_CONFIG_CARRIERFORMAT_SPECIALLIST):
	"""Manages process of extracting payload from carrier."""

	printT("Decoder: Decoder commencing...")

	#Carrier file format validation
	encodedCarrierFileFormat = (encodedCarrierFilePath.split(".")[-1]).lower()
	if(encodedCarrierFileFormat in carrierFormatBlacklist):
		printT("*ERROR* Decoder: Unsupported encoded carrier format detected.")
		return 2, None
	
	#Initialize runtime variable(s)
	encodedCarrierBytes = bytearray()
	encodedCarrierEncodingSignatureBytes = bytearray()

	#Initialize return variable(s)
	decoderMsg = ""

	#Load image type encoded carrier file
	if(encodedCarrierFileFormat in carrierFormatSpecialList.get("images")):
		printT("Encoder: Image carrier format detected.")

		#Negate frame header offset
		carrierHeaderSize = 0

		#Load carrier file
		try:
			encodedCarrierBytes = cv2.imread(encodedCarrierFilePath)
		except FileNotFoundError:
			printT("*ERROR* Decoder: Invalid file path.")
			return 1, None
		except:
			printT("*ERROR* Decoder: Unexpected reading error occured.")
			return 1, None
		
		#Flatten pixel data to single dimensional bytearray
		status, encodedCarrierBytes = multiDimensionArrayFlattener(encodedCarrierBytes)
		if(status == 1):
			printT("*ERROR* Decoder: Unexpected error occured while flattening encoded carrier bytes.")
			return 1, None
		printT("Decoder: Carrier data transformed.")
	#Load word doc type encoded carrier file
	elif(encodedCarrierFileFormat in carrierFormatSpecialList.get("worddoc")):
		printT("Encoder: Word document carrier format detected.")

		#Negate frame header offset
		carrierHeaderSize = 0

		#Load carrier file
		try:
			#Read word document file
			carrierFile = Document(encodedCarrierFilePath)
			#Convert word document text to bytearray
			for paragraph in carrierFile.paragraphs:
				encodedCarrierBytes += bytearray(paragraph.text, "utf-8")
		except FileNotFoundError:
			printT("*ERROR* Encoder: Invalid file path.")
			return 1, None
		except:
			printT("*ERROR* Encoder: Unexpected reading error occured.")
			return 1, None
	#Load non-specific type carrier files
	else:
		printT("Decoder: Loading encoded carrier file...")
		try:
			with open(encodedCarrierFilePath, "rb") as byteReader:
				currentByte = byteReader.read(2048)
				while currentByte:
					encodedCarrierBytes += currentByte
					currentByte = byteReader.read(2048)
		except FileNotFoundError:
			printT("*ERROR* Decoder: Invalid file path.")
			return 1, None
		except:
			printT("*ERROR* Decoder: Unexpected reading error occured.")
			return 1, None
	printT("Decoder: Encoded carrier file loaded.")
	printT("Decoder: Encoded carrier byte count: " + str(len(encodedCarrierBytes)))

	#Read encoding method signature
	if((encodedCarrierFileFormat in carrierFormatSpecialList.get("images")) or
	(encodedCarrierFileFormat in carrierFormatSpecialList.get("worddoc"))):
		encodedCarrierEncodingSignatureBytes = encodedCarrierBytes[:16]
	else:
		encodedCarrierEncodingSignatureBytes = encodedCarrierBytes[carrierHeaderSize:(carrierHeaderSize + 16)]

	#Decode encoding method signature
	status, decodingMethod, encodingParam = payloadFromCarrierDecoder_encodingMethodDetector(
		encodedCarrierEncodingSignatureBytes
	)
	if(status == -1):
		printT("*ERROR* Decoder: Failed to decode encoding method signature.")
		return 1, None
	elif(status == -2):
		printT("*ERROR* Decoder: Unrecognized signature.")
		return 1, None

	#Acquire encoder carrier body
	encodedCarrierBytes = encodedCarrierBytes[(carrierHeaderSize + 16):]
	printT("Decoder: Encoded carrier file decapitated.")
	printT("Decoder: Encoded carrier body byte count: " + str(len(encodedCarrierBytes)))

	#Decode according to specified method
	if(decodingMethod == 0):
		printT("Decoder: Decoding signature of method 0 detected.")
		status, decoderMsg, payloadBytes = twoBitXOrOnionDecoder(
			encodedCarrierBytes)
	elif(decodingMethod == 1):
		printT("Decoder: Decoding signature of method 1 detected.")
		printT("Decoder: Encoding parameter: " + str(encodingParam))
		status, decoderMsg, payloadBytes = bitReplacementDecoder(
			encodedCarrierBytes, encodingParam)

	#When error status is returned, throw to function caller
	if(status == 1):
		printT("*ERROR* Decoder: Decoder method error detected.")
		return 1, decoderMsg
	elif(status == 2):
		printT("*ERROR* Encoder: Carrier file too small.")
		return 2, decoderMsg
	else:
		printT("Decoder: Decoder method successful.")

	#Define payload save path
	decodedPayloadDestPath += "\\" + str(decoderMsg)
	printT("Decoder: Payload save path initialized.")

	#Start writing payload bytes
	printT("Decoder: Writing payload file...")
	try:
		if os.path.exists(decodedPayloadDestPath):
			os.remove(decodedPayloadDestPath)
		with open(decodedPayloadDestPath, "wb") as byteWriter:
			byteWriter.write(payloadBytes)
	except:
		printT("*ERROR* Decoder: Unexpected writing error occured.")
		return 1, None
	printT("Decoder: Writing payload file complete.")
	
	printT("Decoder: Decoder terminating.")

	#Return success status
	return 0, decoderMsg
def payloadFromCarrierDecoder_encodingMethodDetector(
	carrierSignatureBytes):
	"""payloadFromCarrierDecoder's sub function; helps detect 
	encoding method from encoding metadata to auto-select decoding 
	function."""

	#Acquire encoding method signature
	status, methodSignature = twoBitXOrOnionDecoder(
		carrierSignatureBytes, True
	)
	if(status == 1):
		return -1, 0, 0

	#Convert signature bytearray to int array
	methodSignature = [x for x in methodSignature]
	
	carrierSignature = methodSignature[:7]

	#Acquire encoding method through dictionary
	for key in ENCODING_CONFIG_SIGNATURES.keys():
		if(ENCODING_CONFIG_SIGNATURES.get(key) == carrierSignature):
			return 0, key, methodSignature[7]
	return -2, 0, 0

#Method00: "2-bit X-OR Onion" encoder(s) & decoder(s) 
def twoBitXOrOnionEncoder(
	carrierBytes,
	payloadBytes):
	"""Encoding Method 0: Encode 1 payload bit with 2 carrier bits 
	through utilization of X-OR."""

	#Validate carrier size
	requiredCarrierSize = len(payloadBytes) * 2
	if(len(carrierBytes) < requiredCarrierSize):
		#If carrier size is not big enough
		return 2, None, None

	#Initialize snooper variables
	bitsModified = 0

	#Initialize checksum variables
	carrierBytesEncoded = 0; payloadBytesEncoded = 0

	#Initialize return variable
	encodedBytes = carrierBytes.copy()

	#X-Or encoding of 1 payload byte into 2 carrier bytes
	currentCarrierByte = 0; 
	for currentPayloadByte in range(len(payloadBytes)):
		printT("Encoding: " + str(currentPayloadByte) + "/" + str(len(payloadBytes)))

		#Acquire next 16 carrier bits/2 carrier bytes
		carrier16Bits = convertToBinaryStr(carrierBytes[currentCarrierByte])
		carrier16Bits += convertToBinaryStr(carrierBytes[currentCarrierByte + 1])
		carrier16Bits = list(carrier16Bits)

		#Acquire next 8 payload bits
		payload8Bits = convertToBinaryStr(payloadBytes[currentPayloadByte])

		#Encode 1 payload byte into carrier 2 bytes
		payload8BitsSegment = 0; payload8BitsPos = 0
		while(payload8BitsSegment < 2):
			#Pack 4 payload bits into 8 carrier bits
			i = (0 + (payload8BitsSegment * 8)); j = (7 + (payload8BitsSegment * 8))
			while(i < j):
				status, carrier16Bits[j] = twoBitXOrOnionEncoder_XOrComplier( 
					carrier16Bits[i],
					carrier16Bits[j],
					(payload8Bits[payload8BitsPos])
				)
				#Update snooper variable
				if(status == 1): bitsModified += 1
				
				i += 1; j -= 1
				payload8BitsPos += 1
			payload8BitsSegment += 1

		#Transform and store encoded carrier bits as bytes
		carrier16Bits = "".join(carrier16Bits)

		#Move encoded bytes into encoded bytearray
		encodedBytes[currentCarrierByte] = int(carrier16Bits[:8], 2)
		encodedBytes[currentCarrierByte + 1] = int(carrier16Bits[8:], 2)

		#Update Checksum variables
		carrierBytesEncoded += 2
		payloadBytesEncoded += 1
		currentCarrierByte += 2

	#Tally checksum
	if(carrierBytesEncoded != (len(payloadBytes) * 2) or payloadBytesEncoded != len(payloadBytes)):
		#If checksum tally failed
		return 1, None, None

	#Encoding successful
	return 0, encodedBytes, bitsModified
def twoBitXOrOnionEncoder_XOrComplier(
	bit01,
	bit02,
	compliment):
	"""twoBitXOrOnionEncoder's sub function; helps ensure X-OR 
	compliance between 2 bits."""

	#If compliment is 0
	if(compliment == "0"):
		#bit01 != bit02, toggle bit02 to disable X-Or match
		if(bit01 != bit02):
			return 1, bit01

	#If compliment is 1
	else:
		#bit01 == bit02, toggle bit02 to enable X-Or match
		if(bit01 == bit02):
			if(bit02 == "1"):
				return 1, "0"
			else:
				return 1, "1"

	#No bit toggling required
	return 0, bit02
def twoBitXOrOnionEncoder_wordDocErrorCorrection(
	encodedBytes,
	min,
	max):
	"""twoBitXOrOnionEncoder's sub function; helps correct all bytes to 
	comply to X-OR between 2 bits and within decimal value range."""

	for i in range(len(encodedBytes)):
		encodedByteVal = encodedBytes[i]
		if((encodedByteVal < min) or (encodedByteVal > max)):
			binaryStr = convertToBinaryStr(encodedBytes[i])

			print("Decimal Val: " + str(encodedByteVal))
			print("Binary Val: " + str(binaryStr))

			status, binaryStr = twoBitXOrOnionEncoder_wordDocErrorCorrection_XOrComplierWithRangeFactor(binaryStr, min, max)
			
			print("Decimal Val: " + str(binaryStr))
			print("Binary Val: " + str(convertToBinaryStr(binaryStr)))

			if(status == 1):
				return 1, None
			else:
				encodedBytes[i] = binaryStr
	return 0, encodedBytes
def twoBitXOrOnionEncoder_wordDocErrorCorrection_XOrComplierWithRangeFactor(
	encodedByte,
	min,
	max):
	"""twoBitXOrOnionEncoder's sub function; helps correct a byte to 
	comply to X-OR between 2 bits and within decimal value range."""
	
	pairs = [
		[0, 7],
		[1, 6],
		[2, 5],
		[3, 4]
	]

	group = [0, 1, 2, 3]
	combinations = []
	virtualFlippedBin = encodedByte.copy()

	for r in range(len(group) + 1):
		for combination in itertools.combinations(group, r):
			combinations.append(list(combination))
	combinations = combinations[1:]

	for combination in combinations:
		for pairNum in combination:
			virtualFlippedBin[pairs[pairNum][0]], virtualFlippedBin[pairs[pairNum][1]] = twoBitXOrOnionEncoder_wordDocErrorCorrection_toggleBitPairs(
				virtualFlippedBin[pairs[pairNum][0]], virtualFlippedBin[pairs[pairNum][1]])
			
			#Convert str binary to int
			decVal = int("".join(virtualFlippedBin), 2)
			#Check if flipped binary array fufills conditions
			if(decVal >= min and decVal <= max):
				return 0, decVal
			else:
				virtualFlippedBin = encodedByte.copy()
	return 1, None
def twoBitXOrOnionEncoder_wordDocErrorCorrection_toggleBitPairs(
	bit01,
	bit02):
	"""twoBitXOrOnionEncoder's sub function; helps toggle 2 bits."""
	bit01 = "1" if(bit01 == "0") else "0"
	bit02 = "1" if(bit02 == "0") else "0"
	return bit01, bit02
def twoBitXOrOnionDecoder(
	carrierBytes,
	decodingSignature = False):
	"""Decoding Method 0: Decode 1 payload bit with 2 carrier bits 
	through utilization of X-OR."""

	currentCarrierByte = 0
	payloadBytes = bytearray()
	payloadFileNameBytes = bytearray(); payloadFileName = ""

	if not decodingSignature:
		#Extract payload header segment01[1 byte]: Payload file name byte count
		status, extractedByte = twoBitXOrOnionDecoder_encodedByteExtractor(
			carrierBytes[currentCarrierByte],
			carrierBytes[currentCarrierByte + 1]
		)
		if(status == 1):
			return 1, "Unexpected error occurred while extracting payload file name byte count.", None
		currentCarrierByte += 2
		payloadFileNameByteCount = int.from_bytes(extractedByte, 'big')

		printT("Decoder: Payload name byte count: " + str(payloadFileNameByteCount))

		#Extract payload header segment02: Payload file name bytes
		for x in range(payloadFileNameByteCount):
			status, extractedByte = twoBitXOrOnionDecoder_encodedByteExtractor(
				carrierBytes[currentCarrierByte],
				carrierBytes[currentCarrierByte + 1]
			)
			if(status == 1):
				return 1, "Unexpected error occurred.", None
			payloadFileNameBytes += extractedByte
			currentCarrierByte += 2
		try:
			payloadFileName = payloadFileNameBytes.decode("utf-8")
		except:
			return 1, "Unexpected error occurred while decoding payload file name.", None

		#Extract payload header segment03[4 bytes]: Payload body byte count
		payloadBodyByteCountBytes = bytearray()
		for x in range(4):
			status, extractedByte = twoBitXOrOnionDecoder_encodedByteExtractor(
				carrierBytes[currentCarrierByte],
				carrierBytes[currentCarrierByte + 1]
			)
			if(status == 1):
				return 1, "Unexpected error occurred while extracting payload body byte count.", None
			payloadBodyByteCountBytes += extractedByte
			currentCarrierByte += 2
		payloadBodyByteCount = int.from_bytes(payloadBodyByteCountBytes, "big")
		if(payloadBodyByteCount <= 0):
			return 1, "Invalid payload body byte count detected.", None

		printT("Decoder: Payload byte count: " + str(payloadBodyByteCount))
	else:
		payloadBodyByteCount = int((len(carrierBytes))/2)
		currentCarrierByte = 0

	#Extract payload body bytes
	for x in range(payloadBodyByteCount):
		status, extractedByte = twoBitXOrOnionDecoder_encodedByteExtractor(
			carrierBytes[currentCarrierByte],
			carrierBytes[currentCarrierByte + 1]
		)
		if(status == 1):
			return 1, "Unexpected error occurred while extracting payload body.", None
		payloadBytes += extractedByte
		currentCarrierByte += 2
	
	if(decodingSignature):
		return 0, payloadBytes
	else:
		return 0, payloadFileName, payloadBytes
def twoBitXOrOnionDecoder_encodedByteExtractor(
	byte01,
	byte02):
	"""twoBitXOrOnionDecoder's sub function; helps extract 1 payload 
	byte from 2 encoded carrier byte."""

	binList = [
		convertToBinaryStr(byte01),
		convertToBinaryStr(byte02),
	]

	extractedBin = ""
	for x in binList:
		i = 0; j = 7
		while i < j:
			if(x[i] == x[j]):
				extractedBin += "0"
			else:
				extractedBin += "1"
			i += 1; j -= 1

	extractedBin = "".join(extractedBin)
	extractedVal = int(extractedBin, 2)
	extractedByte = extractedVal.to_bytes(1, "big")

	return 0, extractedByte

#Method01: Direct bit replacement encoder(s) & decoder(s) 
def bitReplacementEncoder(
	carrierBytes,
	payloadBytes,
	replacementBitsVal):
	"""Encoding Method 1: Encode payload with direct bit replacement method."""

	#Transforming encoding parameters
	bitReplacementBitCount, replacementBits = bitReplacement_replacementBitIdentifier(
		replacementBitsVal)

	#Validate encoding parameters
	if(replacementBitsVal > 255 or replacementBitsVal < 1):
		return 3, None, None
	elif(bitReplacementBitCount == 0):
		return 3, None, None

	#Validate carrier size
	requiredCarrierByteCount = math.ceil(((len(payloadBytes) * 8)/bitReplacementBitCount))
	if(len(carrierBytes) <= requiredCarrierByteCount):
		return 2, None, None

	#Initialize snooper variables
	bitsModified = 0

	#Initialize runtime variables
	currentBitInPayloadBytePos = 0; currentBitInCarrierBytePos = 0
	currentPayloadBytePos = 0; currentCarrierBytePos = 0
	
	#Initialize return variables
	encodedBytes = carrierBytes.copy()

	#Processing
	currentPayloadByteBinary = list(convertToBinaryStr(payloadBytes[currentPayloadBytePos]))
	currentCarrierByteBinary = list(convertToBinaryStr(carrierBytes[currentCarrierBytePos]))
	wrapItUp = False
	while(currentCarrierBytePos < len(carrierBytes)):
		printT("Encoding: " + str(currentPayloadBytePos) + "/" + str(len(payloadBytes)))

		byteModified = False
		for currentBitInCarrierBytePos in range(len(replacementBits)):
			#If replacement in current carrier bit pos required
			if(replacementBits[currentBitInCarrierBytePos] == 1):
				#If replacement bit not equals to payload bit
				if(currentCarrierByteBinary[currentBitInCarrierBytePos] != currentPayloadByteBinary[currentBitInPayloadBytePos]):
					#Bit replacement
					currentCarrierByteBinary[currentBitInCarrierBytePos] = currentPayloadByteBinary[currentBitInPayloadBytePos]
					#Raise byte modified flag
					byteModified = True
					#Update snooper variable
					bitsModified += 1

				#Check if current payload bit position has reached end of byte
				currentBitInPayloadBytePos += 1
				if(currentBitInPayloadBytePos > 7):
					#Reset payload byte's bit counter for next byte
					currentBitInPayloadBytePos = 0
					#Update payload byte counter to next byte
					currentPayloadBytePos += 1
					#Check if there is next payload byte
					if(currentPayloadBytePos < len(payloadBytes)):
						#Load next payload byte
						currentPayloadByteBinary = list(convertToBinaryStr(payloadBytes[currentPayloadBytePos]))
					else:
						#Raise encoding completion flag
						wrapItUp = True
						break
		
		#Check if carrier byte was modified
		if(byteModified or wrapItUp):
			#Package modified byte into encoded carrier bytes
			encodedBytes[currentCarrierBytePos] = int(("".join(currentCarrierByteBinary)), 2)
			#Check if payload has finished encoding
			if(wrapItUp):
				break
		
		#Load next carrier byte for encoding
		currentCarrierBytePos += 1
		try:
			currentCarrierByteBinary = list(convertToBinaryStr(carrierBytes[currentCarrierBytePos]))
		except IndexError:
			#Not enough carrier bytes to encode payload bytes in
			return 2, None, None

	#Tally checksum
	if((currentPayloadBytePos) != len(payloadBytes)):
		#If checksum tally failed
		return 1, None, None

	return 0, encodedBytes, bitsModified
def bitReplacementDecoder(
	carrierBytes,
	replacementBits):
	"""Decoding Method 1: Decode payload with direct bit replacement method."""

	#Initialize runtime variables
	currentBitInPayloadBytePos = 0; currentBitInCarrierBytePos = 0
	currentPayloadBytePos = 0; currentCarrierBytePos = 0
	checkpointVal = 1; checkpointCount = 0

	#Initialize return variables
	payloadBytes = bytearray()
	payloadFilename = ""

	#Transforming encoding parameters
	bitReplacementBitCount, replacementBits = bitReplacement_replacementBitIdentifier(
		replacementBits)

	currentCarrierByteBinary = list(convertToBinaryStr(carrierBytes[currentCarrierBytePos]))
	payloadConstructed = False
	payloadByteBinary = []
	for currentCarrierBytePos in range(len(carrierBytes)):
		for currentBitInCarrierBytePos in range(len(replacementBits)):
			#Identify payload bits according to replacementBits param
			if(replacementBits[currentBitInCarrierBytePos] == 1):
				#Extract bit to payload
				payloadByteBinary.append(currentCarrierByteBinary[currentBitInCarrierBytePos])
				
				#Check if current payload bit position has reached end of byte
				currentBitInPayloadBytePos += 1
				if(currentBitInPayloadBytePos > 7):
					#Pack binary into byte
					payloadBytes += (int(("".join(payloadByteBinary)), 2)).to_bytes(1, "big")
					payloadByteBinary.clear()

					#Reset payload byte's bit counter for next byte
					currentBitInPayloadBytePos = 0
					#Update payload byte counter to next byte
					currentPayloadBytePos += 1

					#Check for any notable byte segments have been decoded
					#Yes, it looks inefficient, because I am tired
					if(currentPayloadBytePos == checkpointVal):
						#Segment01[1 byte] of payload bytes found, byte count of payload filename
						if(checkpointCount == 0):
							#Acquire payload filename byte count
							payloadFileNameByteCount = payloadBytes[0]

							#Set next check point at end of payload filename bytes
							checkpointVal += payloadFileNameByteCount
							#Ready next checkpoint flag
							checkpointCount += 1
							#Clear payloadBytes
							payloadBytes.clear()

						#Segment02 of payload bytes found, bytes of payload filename
						elif(checkpointCount == 1):
							payloadFilename = payloadBytes.decode("utf-8")

							#Set next check point at end of payload byte count bytes
							checkpointVal += 4
							#Ready next checkpoint flag
							checkpointCount += 1
							#Clear payloadBytes
							payloadBytes.clear()

						#Segment03[4 bytes] of payload bytes found, byte count of payload bytes
						elif(checkpointCount == 2):
							#Acquire payload byte count
							payloadByteCount = int.from_bytes(payloadBytes, "big")

							#Set next check point at end of payload bytes
							checkpointVal += payloadByteCount
							#Ready next checkpoint flag
							checkpointCount += 1
							#Clear payloadBytes
							payloadBytes.clear()

						#End of payload found
						elif(checkpointCount == 3):
							#Exit carrier bytes iterating loop
							payloadConstructed = True
							break;

		if payloadConstructed:
			return 0, payloadFilename, payloadBytes
		else:
			#Load next carrier byte for encoding
			currentCarrierBytePos += 1
			currentCarrierByteBinary = list(convertToBinaryStr(carrierBytes[currentCarrierBytePos]))

	#Should not reach here
	return 1, None, None
def bitReplacement_replacementBitIdentifier(
	encodingParam):
	"""Direct bit replacement encoding method's sub function; helps decode 
	decimal value to binary to identify which bits will be used for 
	replacement."""

	binaryStr = convertToBinaryStr(encodingParam)
	binaryIntList = list(map(lambda x: int(x), list(binaryStr)))
	bitCount = sum(map(lambda x : x == 1, binaryIntList))
	
	return bitCount, binaryIntList

#Method02[Implicit]: Word document bit replacement encoder(s) & decoder(s)
def bitReplacementEncoder_wordDoc(
	carrierBytes,
	payloadBytes,
	replacementBitsVal,
	reservedCharDecVal = ENCODING_BITREPLACEMENT_WORDDOC_RESERVED):
	"""Encoding Method 2: Encode doc payload with direct bit replacement method."""

	#Transforming encoding parameters
	bitReplacementBitCount, replacementBits = bitReplacement_replacementBitIdentifier(
		replacementBitsVal)

	#Validate encoding parameters
	if(replacementBitsVal > 255 or replacementBitsVal < 1):
		return 3, None, None
	elif(bitReplacementBitCount == 0):
		return 3, None, None

	#Initialize snooper variables
	bitsModified = 0

	#Initialize runtime variables
	currentCarrierBytePos = 0;
	currentTruePayloadBytePos = 0; currentVirtualPayloadBytePos = 0
	currentTrueBitInPayloadBytePos = 0; currentVirtualBitInPayloadBytePos = 0

	#Initialize return variables
	encodedBytes = carrierBytes.copy()

	#Processing
	currentPayloadByteBinary = list(convertToBinaryStr(payloadBytes[currentVirtualPayloadBytePos]))
	currentCarrierByteBinary = list(convertToBinaryStr(carrierBytes[currentCarrierBytePos]))
	wrapItUp = False
	while(currentCarrierBytePos < len(carrierBytes)):
		byteModified = False
		for currentBitInCarrierBytePos in range(len(replacementBits)):
			#If replacement in current carrier bit pos required
			if(replacementBits[currentBitInCarrierBytePos] == 1):
				#If replacement bit not equals to payload bit
				if(currentCarrierByteBinary[currentBitInCarrierBytePos] != currentPayloadByteBinary[currentVirtualBitInPayloadBytePos]):
					#Bit replacement
					currentCarrierByteBinary[currentBitInCarrierBytePos] = currentPayloadByteBinary[currentVirtualBitInPayloadBytePos]
					#Raise byte modified flag
					byteModified = True
					#Update snooper variable
					bitsModified += 1

				#Check if current payload bit position has reached end of byte
				currentVirtualBitInPayloadBytePos += 1
				if(currentVirtualBitInPayloadBytePos > 7):
					#Reset payload byte's bit counter for next byte
					currentVirtualBitInPayloadBytePos = 0
					#Update payload byte counter to next byte
					currentVirtualPayloadBytePos += 1
					#Check if there is next payload byte
					if(currentVirtualPayloadBytePos < len(payloadBytes)):
						#Load next payload byte
						currentPayloadByteBinary = list(convertToBinaryStr(payloadBytes[currentVirtualPayloadBytePos]))
					else:
						#Raise encoding completion flag
						wrapItUp = True
						break
		
		#Check if carrier byte was modified
		if(byteModified or wrapItUp):
			#Package modified byte into encoded carrier bytes
			encodedByteVal = int(("".join(currentCarrierByteBinary)), 2)
			#Check if encoded byte becomes a reserved or illegal character
			#Range of 32 - 126 defined by UTF-8 legal character set
			if((encodedByteVal < 32) or (encodedByteVal > 126) or (encodedByteVal == reservedCharDecVal)):
				#Change carrier byte to reserved char as it is unable to carry payload
				encodedBytes[currentCarrierBytePos] = reservedCharDecVal
				#Reload previous payload byte if required
				if(currentTruePayloadBytePos != currentVirtualPayloadBytePos):
					currentPayloadByteBinary = list(convertToBinaryStr(payloadBytes[currentTruePayloadBytePos]))
				#Revert virtual counters actual counter
				currentVirtualPayloadBytePos = currentTruePayloadBytePos
				currentVirtualBitInPayloadBytePos = currentTrueBitInPayloadBytePos
			#Encoded byte is valid
			else:
				#Update encoded byte into encoded carrier
				encodedBytes[currentCarrierBytePos] = encodedByteVal
				#Update virtual counter to actual counter
				currentTruePayloadBytePos = currentVirtualPayloadBytePos
				currentTrueBitInPayloadBytePos = currentVirtualBitInPayloadBytePos

			#Check if payload has finished encoding
			if(wrapItUp):
				#Tally checksum
				if((currentVirtualPayloadBytePos) != len(payloadBytes)):
					#If checksum tally failed
					return 1, None, None
				break
		
		#Load next carrier byte for encoding
		currentCarrierBytePos += 1
		try:
			currentCarrierByteBinary = list(convertToBinaryStr(carrierBytes[currentCarrierBytePos]))
		except IndexError:
			return 2, None, None

	return 0, encodedBytes, bitsModified
def bitReplacementDecoder_wordDoc(
	carrierBytes,
	replacementBitsVal,
	reservedCharDecVal = ENCODING_BITREPLACEMENT_WORDDOC_RESERVED):
	"""Decoding Method 2: Decode doc payload with direct bit replacement method."""

	#Transforming encoding parameters
	bitReplacementBitCount, replacementBits = bitReplacement_replacementBitIdentifier(
		replacementBitsVal)

	#Validate encoding parameters
	if(replacementBitsVal > 255 or replacementBitsVal < 1):
		return 3, None, None
	elif(bitReplacementBitCount == 0):
		return 3, None, None

	#Initialize runtime variables
	currentBitInCarrierBytePos = 0; currentCarrierBytePos = 0
	currentBitInPayloadBytePos = 0
	currentPayloadBytePos = 0
	checkpointVal = 1; checkpointCount = 0

	#Initialize return variables
	payloadBytes = bytearray()
	payloadFilename = ""

	#Processing
	currentCarrierByteBinary = list(convertToBinaryStr(carrierBytes[currentCarrierBytePos]))
	payloadConstructed = False
	payloadByteBinary = []
	for currentCarrierBytePos in range(len(carrierBytes)):
		currentByteVal = carrierBytes[currentCarrierBytePos]
		if(currentByteVal < 32) or (currentByteVal > 126) or (currentByteVal == reservedCharDecVal):
			pass
		else:
			for currentBitInCarrierBytePos in range(len(replacementBits)):
				#Identify payload bits according to replacementBits param
				if(replacementBits[currentBitInCarrierBytePos] == 1):
					#Extract bit to payload
					payloadByteBinary.append(currentCarrierByteBinary[currentBitInCarrierBytePos])
					
					#Check if current payload bit position has reached end of byte
					currentBitInPayloadBytePos += 1
					if(currentBitInPayloadBytePos > 7):
						#Pack binary into byte
						payloadBytes += (int(("".join(payloadByteBinary)), 2)).to_bytes(1, "big")
						payloadByteBinary.clear()

						#Reset payload byte's bit counter for next byte
						currentBitInPayloadBytePos = 0
						#Update payload byte counter to next byte
						currentPayloadBytePos += 1

						#Check for any notable byte segments have been decoded
						#Yes, it looks inefficient, because I am tired
						if(currentPayloadBytePos == checkpointVal):
							#Segment01[1 byte] of payload bytes found, byte count of payload filename
							if(checkpointCount == 0):
								#Acquire payload filename byte count
								payloadFileNameByteCount = payloadBytes[0]

								#Set next check point at end of payload filename bytes
								checkpointVal += payloadFileNameByteCount
								#Ready next checkpoint flag
								checkpointCount += 1
								#Clear payloadBytes
								payloadBytes.clear()

							#Segment02 of payload bytes found, bytes of payload filename
							elif(checkpointCount == 1):
								payloadFilename = payloadBytes.decode("utf-8")

								#Set next check point at end of payload byte count bytes
								checkpointVal += 4
								#Ready next checkpoint flag
								checkpointCount += 1
								#Clear payloadBytes
								payloadBytes.clear()

							#Segment03[4 bytes] of payload bytes found, byte count of payload bytes
							elif(checkpointCount == 2):
								#Acquire payload byte count
								payloadByteCount = int.from_bytes(payloadBytes, "big")

								#Set next check point at end of payload bytes
								checkpointVal += payloadByteCount
								#Ready next checkpoint flag
								checkpointCount += 1
								#Clear payloadBytes
								payloadBytes.clear()

							#End of payload found
							elif(checkpointCount == 3):
								#Exit carrier bytes iterating loop
								payloadConstructed = True
								break;

		if payloadConstructed:
			return 0, payloadFilename, payloadBytes
		else:
			#Load next carrier byte for encoding
			currentCarrierBytePos += 1
			currentCarrierByteBinary = list(convertToBinaryStr(carrierBytes[currentCarrierBytePos]))

	#Should not reach here
	return 1, None, None

#Helper function(s)
def multiDimensionArrayFlattener(
	inputArray):
	"""Flattens multi dimension array into single dimension."""
	try:
		returnVal = (np.array(inputArray)).flatten().tolist()
		return 0, returnVal
	except:
		return 1, None
def multiDimensionArrayReshaper(
	inputArray,
	inputDimensions):
	"""Shapes a single dimension array to specified dimensions."""
	try:
		returnVal = np.reshape(inputArray, inputDimensions).tolist()
		return 0, returnVal
	except:
		return 1, None
def convertToBinaryStr(
	input):
	"""Converts byte/int to binary in form of a string."""
	return list((bin(input))[2:].zfill(8))
def printT(
	inputString,
	nextLine = False,
	consoleOutput = PROGRAM_CONFIG_CONSOLEOUTPUT):
	"""Custom print function that prepends current timestamp to string 
	when printing."""

	if(consoleOutput == 0): return

	if nextLine: print()
	ts = datetime.now().strftime("%H:%M:%S.%f")[:-4]
	print(ts + "\t" + str(inputString))

main()